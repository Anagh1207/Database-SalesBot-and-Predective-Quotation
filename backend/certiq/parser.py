"""
Universal Document Parser — Layer 1

Handles ANY input format:
- PDF (PyMuPDF — fast, handles scanned via OCR)
- Word (.docx)
- Plain text / email text
- Form submission (dict)

Returns a standardised ParsedDocument object that the
attribute extractor can work with regardless of input type.
"""

import re
import fitz          # PyMuPDF
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional
from app.logger import get_logger

logger = get_logger("parser")


@dataclass
class ParsedDocument:
    """
    Standardised output from the parser.
    The attribute extractor only ever sees this — never the raw file.
    """
    source_type: str          # pdf / docx / text / form
    source_path: str          # file path or "inline"
    raw_text: str             # full extracted text
    text_lower: str           # lowercase version for keyword matching
    sections: List[Dict]      # list of {heading, content} blocks
    metadata: Dict[str, Any]  # cert_no, company, pages, etc.
    word_count: int
    page_count: int


def parse_pdf(pdf_path: str) -> ParsedDocument:
    """
    Extracts text from a PDF using PyMuPDF.
    Handles both text-based and scanned PDFs.
    Detects document sections by font size (headings are larger).
    """
    path = Path(pdf_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    pdf_doc = fitz.open(pdf_path)  # ✅ renamed to pdf_doc to avoid collision
    full_text = ""
    sections = []
    current_section = {"heading": "", "content": ""}

    try:
        for page_num, page in enumerate(pdf_doc):
            # Extract text with formatting info
            blocks = page.get_text("dict")["blocks"]

            for block in blocks:
                if block.get("type") != 0:  # 0 = text block
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        size = span.get("size", 10)
                        if not text:
                            continue

                        full_text += text + " "

                        # Detect headings by font size (headings > 11pt typically)
                        if size > 11 and len(text) > 3 and len(text) < 80:
                            # Save current section
                            if current_section["content"].strip():
                                sections.append(current_section.copy())
                            current_section = {"heading": text, "content": ""}
                        else:
                            current_section["content"] += text + " "

        # Save last section
        if current_section["content"].strip():
            sections.append(current_section)

        # ✅ Read page_count BEFORE closing the document
        page_count = len(pdf_doc)

    finally:
        pdf_doc.close()  # ✅ always close, even if an error occurs

    # Extract metadata from text
    metadata = _extract_pdf_metadata(full_text)
    metadata["pages"] = page_count   # ✅ use the captured value
    metadata["filename"] = path.name

    return ParsedDocument(
        source_type="pdf",
        source_path=str(pdf_path),
        raw_text=full_text,
        text_lower=full_text.lower(),
        sections=sections,
        metadata=metadata,
        word_count=len(full_text.split()),
        page_count=page_count,        # ✅ use the captured value
    )


def parse_docx(docx_path: str) -> ParsedDocument:
    """
    Extracts text from a Word document.
    Preserves heading structure.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    path = Path(docx_path)
    doc = Document(docx_path)

    full_text = ""
    sections = []
    current_section = {"heading": "", "content": ""}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        full_text += text + " "

        # Detect headings by style name
        if para.style.name.startswith("Heading"):
            if current_section["content"].strip():
                sections.append(current_section.copy())
            current_section = {"heading": text, "content": ""}
        else:
            current_section["content"] += text + " "

    if current_section["content"].strip():
        sections.append(current_section)

    metadata = _extract_pdf_metadata(full_text)
    metadata["filename"] = path.name

    return ParsedDocument(
        source_type="docx",
        source_path=str(docx_path),
        raw_text=full_text,
        text_lower=full_text.lower(),
        sections=sections,
        metadata=metadata,
        word_count=len(full_text.split()),
        page_count=0,
    )


def parse_text(text: str, source_label: str = "inline") -> ParsedDocument:
    """
    Parses plain text — from email, chatbot input, copy-paste.
    Splits into sections on double newlines or numbered headings.
    """
    full_text = text
    sections = []

    # Split on paragraph breaks
    paragraphs = re.split(r"\n{2,}", text)
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        # Simple heading detection — all caps or ends with colon
        if re.match(r"^[A-Z][A-Z\s]{3,}:?$", para) or re.match(r"^\d+\.", para):
            sections.append({"heading": para, "content": ""})
        else:
            if sections:
                sections[-1]["content"] += para + " "
            else:
                sections.append({"heading": "", "content": para})

    metadata = _extract_pdf_metadata(full_text)
    metadata["filename"] = source_label

    return ParsedDocument(
        source_type="text",
        source_path=source_label,
        raw_text=full_text,
        text_lower=full_text.lower(),
        sections=sections,
        metadata=metadata,
        word_count=len(full_text.split()),
        page_count=0,
    )


def parse_form(form_data: Dict[str, Any]) -> ParsedDocument:
    """
    Converts a filled-in form (dict) into a ParsedDocument.
    Used when user fills in the dynamic form in the UI.
    """
    # Convert form fields to text
    lines = []
    for key, value in form_data.items():
        lines.append(f"{key}: {value}")
    full_text = "\n".join(lines)

    return ParsedDocument(
        source_type="form",
        source_path="form_submission",
        raw_text=full_text,
        text_lower=full_text.lower(),
        sections=[{"heading": "Form Submission", "content": full_text}],
        metadata={"filename": "form_submission", **form_data},
        word_count=len(full_text.split()),
        page_count=0,
    )


def parse_any(input_data: Any) -> ParsedDocument:
    """
    Universal entry point — auto-detects input type and parses it.

    Accepts:
    - str path to PDF → parse_pdf
    - str path to DOCX → parse_docx
    - str text → parse_text
    - dict → parse_form
    """
    if isinstance(input_data, dict):
        return parse_form(input_data)

    if isinstance(input_data, str):
        path = Path(input_data)

        if path.exists():
            suffix = path.suffix.lower()
            if suffix == ".pdf":
                return parse_pdf(input_data)
            elif suffix in [".docx", ".doc"]:
                return parse_docx(input_data)
            else:
                return parse_text(path.read_text(encoding="utf-8"), str(path))
        else:
            # Treat as raw text
            return parse_text(input_data)

    raise ValueError(f"Unsupported input type: {type(input_data)}")


def _extract_pdf_metadata(text: str) -> Dict[str, Any]:
    """Extracts common metadata fields from document text."""
    text_lower = text.lower()

    # Cert number
    cert_match = re.search(r"(\d{2}/\d{4,5})", text)
    cert_no = cert_match.group(1) if cert_match else ""

    # BBA cert number format
    bba_match = re.search(r"bba\s+(\d{2}/\d{4,5})", text_lower)
    if bba_match:
        cert_no = bba_match.group(1)

    # Company name — usually in first 10 lines
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    company = ""
    for line in lines[:15]:
        if len(line) > 5 and len(line) < 80:
            if not re.match(r"^page \d+", line.lower()):
                if not re.match(r"^\d", line):
                    company = line
                    break

    return {
        "cert_no": cert_no,
        "company": company,
    }